from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from deep_translator import GoogleTranslator
from io import BytesIO
from docx import Document
import base64
import os

app = Flask(__name__)

# More robust CORS configuration
CORS(app, resources={r"/upload": {"origins": "*"}}, 
     methods=["POST", "OPTIONS"], 
     allow_headers=["Content-Type", "Authorization"])

def allowed_file(file_data):
    """Check if the uploaded file is of a valid type (e.g., docx)."""
    return file_data[:2] == b"PK"  # DOCX files start with "PK"

def extract_content(docx_file_data):
    """Extract text and table content from the DOCX file."""
    file_stream = BytesIO(docx_file_data)
    doc = Document(file_stream)
    content = {"text": [], "tables": []}

    # Extract paragraphs
    for para in doc.paragraphs:
        content["text"].append(para.text)

    # Extract table data
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append(table_data)

    return doc, content

def translate_text_in_place(doc, target_language):
    """Translate text directly within the DOCX document."""
    translator = GoogleTranslator(source='auto', target=target_language)

    # Translate paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                try:
                    run.text = translator.translate(run.text)
                except Exception as e:
                    print(f"Error translating paragraph text: {run.text}. Error: {e}")

    # Translate table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            try:
                                run.text = translator.translate(run.text)
                            except Exception as e:
                                print(f"Error translating table cell text: {run.text}. Error: {e}")

    return doc

@app.route('/upload', methods=['POST', 'OPTIONS'])
def upload_file():
    """Handle file upload, translation, and return the translated file."""
    # Handle CORS preflight request
    if request.method == 'OPTIONS':
        response = jsonify({})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST')
        return response

    try:
        # Parse JSON payload
        data = request.get_json()

        # Decode the base64 file data
        file_data = base64.b64decode(data['file'].split(",")[1])

        # Validate the uploaded file type
        if not allowed_file(file_data):
            return jsonify({'error': 'Invalid file type. Only DOCX files are supported.'}), 400

        # Extract and translate content
        doc, content = extract_content(file_data)
        target_language = data.get('language', 'en').strip()
        if not target_language:
            return jsonify({'error': 'Invalid target language specified.'}), 400

        translated_doc = translate_text_in_place(doc, target_language)

        # Prepare the translated document as an in-memory byte stream
        byte_stream = BytesIO()
        translated_doc.save(byte_stream)
        byte_stream.seek(0)

        # Encode the translated document into base64
        file_base64 = base64.b64encode(byte_stream.read()).decode('utf-8')
        
        # Create response with CORS headers
        response = jsonify({
            'file_url': f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{file_base64}'
        })
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        return response, 200

    except Exception as e:
        # Create error response with CORS headers
        response = jsonify({'error': f"An error occurred: {str(e)}"})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        return response, 500

if __name__ == '__main__':
    app.run(debug=True, port=5001)